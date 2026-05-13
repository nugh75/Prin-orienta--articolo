#!/usr/bin/env python3
"""
Strumento per generare documenti .docx conformi alle norme redazionali
Roma TrE-PRESS, basato sul template `template_Psicology & Education.docx`.

Caratteristiche:
- Template-based: preserva header, footer, impostazioni di pagina
- Formattazione automatica secondo le norme (Times New Roman, interlinea esatta 13pt, etc.)
- Supporto per capitoli, paragrafi, sottoparagrafi, didascalie, bibliografia

Uso come modulo:
    from docx_renderer import TrEPRESSRenderer
    doc = TrEPRESSRenderer("raw/articles/template_Psicology & Education.docx")
    doc.add_chapter("1", "Titolo Capitolo", "Mario Rossi")
    doc.add_section("1", "Titolo Paragrafo")
    doc.add_paragraph("Testo del paragrafo...")
    doc.save("output.docx")

Uso da CLI:
    python scripts/docx_renderer.py input.md output.docx
    python scripts/docx_renderer.py input.md output.docx --template percorso/template.docx
"""

import re
import sys
import argparse
from docx import Document
from docx.shared import Pt, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ── Costanti dal template ──────────────────────────────────────────────

FONT_BODY = "Times New Roman"
FONT_ACCENT = "Tahoma"  # usato solo per caratteri di controllo nel template

SIZE_CHAPTER = Pt(14)       # 28 half-points
SIZE_BODY = Pt(12)          # 24 half-points (default doc)
SIZE_CAPTION = Pt(9)        # 18 half-points
SIZE_BIBLIO = Pt(12)

LINE_SPACING_TWIPS = 260    # esatto 13pt (dal template: lineRule=exact)

INDENT_FIRST_LINE_TWIPS = 284   # ~0.5 cm per capoverso
INDENT_LEFT_BIB_TWIPS = 284     # rientro sinistro per bibliografia (hanging)

ALIGN_CENTER = WD_ALIGN_PARAGRAPH.CENTER
ALIGN_JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY
ALIGN_LEFT = WD_ALIGN_PARAGRAPH.LEFT


# ── Helper XML ─────────────────────────────────────────────────────────

def _set_spacing(paragraph, line=260, line_rule="exact", before=None, after=None):
    """Set paragraph spacing at XML level for full control."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = paragraph._element.makeelement(qn('w:pPr'), {})
        paragraph._element.insert(0, pPr)

    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = pPr.makeelement(qn('w:spacing'), {})
        pPr.append(spacing)

    spacing.set(qn('w:line'), str(line))
    spacing.set(qn('w:lineRule'), line_rule)
    if before is not None:
        spacing.set(qn('w:before'), str(before))
    if after is not None:
        spacing.set(qn('w:after'), str(after))


def _set_indent(paragraph, first_line=None, left=None, hanging=None):
    """Set paragraph indent at XML level."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = paragraph._element.makeelement(qn('w:pPr'), {})
        paragraph._element.insert(0, pPr)

    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = pPr.makeelement(qn('w:ind'), {})
        pPr.append(ind)

    for attr, val in [('firstLine', first_line), ('left', left), ('hanging', hanging)]:
        if val is not None:
            ind.set(qn('w:' + attr), str(val))


def _set_alignment(paragraph, align):
    """Set paragraph alignment."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = paragraph._element.makeelement(qn('w:pPr'), {})
        paragraph._element.insert(0, pPr)

    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = pPr.makeelement(qn('w:jc'), {})
        pPr.append(jc)

    align_map = {
        ALIGN_CENTER: 'center',
        ALIGN_JUSTIFY: 'both',
        ALIGN_LEFT: 'left',
    }
    jc.set(qn('w:val'), align_map.get(align, 'both'))


def _add_run(paragraph, text, font_name=FONT_BODY, size=None, italic=False, bold=False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    # Imposta proprietà solo se True, per evitare w:i w:val="0" inutili
    if italic:
        run.italic = True
    if bold:
        run.bold = True
    # Ensure Times New Roman also for east-asia / complex script
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = run._element.makeelement(qn('w:rPr'), {})
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    return run


def _new_paragraph(doc, alignment=ALIGN_JUSTIFY):
    """Add a new empty paragraph and set alignment."""
    p = doc.add_paragraph()
    _set_alignment(p, alignment)
    _set_spacing(p)
    return p


def _clear_body(doc):
    """Remove all paragraphs from the body, preserving header/footer/sections."""
    body = doc.element.body
    # Remove all w:p elements from body
    for p in body.findall(qn('w:p')):
        body.remove(p)


# ── Renderer principale ──────────────────────────────────────────────

class TrEPRESSRenderer:
    """Renderizza contenuti accademici in .docx secondo le norme Roma TrE-PRESS."""

    def __init__(self, template_path):
        self.doc = Document(template_path)
        _clear_body(self.doc)

    # ── Metodi pubblici ────────────────────────────────────────────────

    def add_chapter(self, chapter_num, title, author=None):
        """Aggiunge blocco capitolo: CAPITOLO X + titolo + autore.

        Args:
            chapter_num: Numero del capitolo (int o str)
            title: Titolo del capitolo
            author: Nome autore (opzionale, in corsivo centrato)
        """
        self.add_chapter_number(chapter_num)
        self.add_chapter_title(title)
        if author:
            self.add_author(author)

    def add_chapter_number(self, number):
        """Capitolo X, centrato, 14pt."""
        p = _new_paragraph(self.doc, ALIGN_CENTER)
        _add_run(p, f"Capitolo {number}", size=SIZE_CHAPTER)
        return p

    def add_chapter_title(self, title):
        """Titolo capitolo, centrato, 14pt."""
        p = _new_paragraph(self.doc, ALIGN_CENTER)
        _add_run(p, title, size=SIZE_CHAPTER)
        return p

    def add_author(self, name):
        """Nome autore, centrato, corsivo, 12pt."""
        p = _new_paragraph(self.doc, ALIGN_CENTER)
        _add_run(p, name, size=SIZE_BODY, italic=True)
        return p

    def add_section_heading(self, title):
        """Sezione senza numero, centrata, 12pt."""
        p = _new_paragraph(self.doc, ALIGN_CENTER)
        _add_run(p, title, size=SIZE_BODY)
        return p

    def add_section(self, number, title):
        """Sezione: '1. Titolo', centrata, 12pt."""
        p = _new_paragraph(self.doc, ALIGN_CENTER)
        _add_run(p, f"{number}. {title}", size=SIZE_BODY)
        return p

    def add_subsection(self, number, title):
        """Sottoparagrafo: '1.1 Titolo', giustificato, corsivo."""
        p = _new_paragraph(self.doc, ALIGN_JUSTIFY)
        _set_indent(p, first_line=INDENT_FIRST_LINE_TWIPS)
        _add_run(p, f"{number} {title}", size=SIZE_BODY, italic=True)
        return p

    def add_subsubsection(self, number, title):
        """Sotto-sottoparagrafo: '1.1.1 Titolo', giustificato, corsivo."""
        p = _new_paragraph(self.doc, ALIGN_JUSTIFY)
        _set_indent(p, first_line=INDENT_FIRST_LINE_TWIPS)
        _add_run(p, f"{number} {title}", size=SIZE_BODY, italic=True)
        return p

    def add_paragraph(self, text):
        """Corpo del testo, giustificato, rientro prima riga, Times 12.

        Supporta formattazione inline:
        - *testo* → corsivo
        - **testo** → grassetto (sconsigliato dalle norme ma gestito)
        """
        p = _new_paragraph(self.doc, ALIGN_JUSTIFY)
        _set_indent(p, first_line=INDENT_FIRST_LINE_TWIPS)
        self._add_rich_text(p, text, SIZE_BODY)
        return p

    def add_paragraph_no_indent(self, text):
        """Come add_paragraph ma senza rientro (es. dopo titolo sezione)."""
        p = _new_paragraph(self.doc, ALIGN_JUSTIFY)
        self._add_rich_text(p, text, SIZE_BODY)
        return p

    def add_figure_caption(self, text):
        """Didascalia figura/tabella, centrata, 9pt."""
        p = _new_paragraph(self.doc, ALIGN_CENTER)
        _add_run(p, text, size=SIZE_CAPTION)
        return p

    def add_empty_line(self):
        """Aggiunge riga vuota."""
        p = _new_paragraph(self.doc, ALIGN_LEFT)
        return p

    def add_abstract(self, text, lang="it"):
        """Abstract, giustificato, 12pt.

        Args:
            text: Testo dell'abstract (supporta *corsivo* e **grassetto**)
            lang: "it" o "en"
        """
        label = "Abstract" if lang == "en" else "Abstract"
        p = _new_paragraph(self.doc, ALIGN_JUSTIFY)
        _add_run(p, f"{label}. ", size=SIZE_BODY, italic=True)
        self._add_rich_text(p, text, SIZE_BODY)
        return p

    def add_keywords(self, keywords, lang="it"):
        """Parole chiave, giustificato.

        Args:
            keywords: Lista di stringhe
            lang: "it" o "en"
        """
        label = "Keywords" if lang == "en" else "Parole chiave"
        p = _new_paragraph(self.doc, ALIGN_JUSTIFY)
        _add_run(p, f"{label}: ", size=SIZE_BODY, italic=True)
        _add_run(p, ", ".join(keywords), size=SIZE_BODY)
        return p

    def add_bibliography_title(self):
        """Titolo 'Bibliografia', centrato, 14pt."""
        p = _new_paragraph(self.doc, ALIGN_CENTER)
        _add_run(p, "Bibliografia", size=SIZE_CHAPTER)
        return p

    def add_bibliography_entry(self, text, title_start=None, title_end=None):
        """Voce bibliografica, giustificata, con rientro sporgente.

        La porzione di titolo (tra title_start e title_end) viene resa in corsivo.

        Args:
            text: Testo completo della voce
            title_start: Indice di inizio del titolo (opzionale)
            title_end: Indice di fine del titolo (opzionale)
        """
        p = _new_paragraph(self.doc, ALIGN_JUSTIFY)
        _set_indent(p, left=INDENT_LEFT_BIB_TWIPS, hanging=INDENT_FIRST_LINE_TWIPS)

        if title_start is not None and title_end is not None:
            # Parte prima del titolo
            _add_run(p, text[:title_start], size=SIZE_BIBLIO)
            # Titolo in corsivo
            _add_run(p, text[title_start:title_end], size=SIZE_BIBLIO, italic=True)
            # Resto
            _add_run(p, text[title_end:], size=SIZE_BIBLIO)
        else:
            _add_run(p, text, size=SIZE_BIBLIO)
        return p

    def add_bibliography_entry_parts(self, *parts):
        """Voce bibliografica con parti alternate tondo/corsivo.

        Args:
            *parts: Sequenza di tuple (text, italic_bool)
                    Es: ("M. ROSSI, ", False), ("Titolo libro", True), (", Editore, 2020.", False)
        """
        p = _new_paragraph(self.doc, ALIGN_JUSTIFY)
        _set_indent(p, left=INDENT_LEFT_BIB_TWIPS, hanging=INDENT_FIRST_LINE_TWIPS)
        for text, italic in parts:
            _add_run(p, text, size=SIZE_BIBLIO, italic=italic)
        return p

    def add_raw_paragraph(self, text, alignment=ALIGN_JUSTIFY, font_size=SIZE_BODY,
                          italic=False, first_line_indent=None, left_indent=None):
        """Paragrafo raw con controllo totale."""
        p = _new_paragraph(self.doc, alignment)
        if first_line_indent is not None:
            _set_indent(p, first_line=first_line_indent)
        if left_indent is not None:
            _set_indent(p, left=left_indent)
        _add_run(p, text, size=font_size, italic=italic)
        return p

    def save(self, output_path):
        """Salva il documento."""
        self.doc.save(output_path)

    # ── Interni ────────────────────────────────────────────────────────

    def _add_rich_text(self, paragraph, text, base_size):
        """Interpreta *italic* e **bold** nel testo."""
        pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)'
        parts = re.split(pattern, text)
        for part in parts:
            if part.startswith('***') and part.endswith('***'):
                _add_run(paragraph, part[3:-3], size=base_size, italic=True, bold=True)
            elif part.startswith('**') and part.endswith('**'):
                _add_run(paragraph, part[2:-2], size=base_size, bold=True)
            elif part.startswith('*') and part.endswith('*'):
                _add_run(paragraph, part[1:-1], size=base_size, italic=True)
            else:
                _add_run(paragraph, part, size=base_size)


# ── CLI: convertitore Markdown → DOCX ─────────────────────────────────

def markdown_to_docx(md_content, template_path):
    """Converte markdown strutturato in docx formattato.

    Convenzioni markdown:
    # CAPITOLO X          → CAPITOLO X (centrato, 14pt)
    ## Titolo Capitolo    → Titolo capitolo (sotto)
    ### Nome Autore       → Autore (corsivo centrato)
    #### 1. Paragrafo     → Sezione (giustificato, rientro)
    #### 1.1 Sottopar.    → Sottosezione (corsivo)
    #### 1.1.1 Sottosottop. → Sotto-sottosezione (corsivo)
    **Figura X – Didascalia** → Didascalia (centrata, 9pt)
    ---                  → Titolo Bibliografia
    - Voce bibliografica → Voce con rientro sporgente
                            (corsivo tra parentesi quadre [titolo])
    """
    renderer = TrEPRESSRenderer(template_path)
    in_bibliography = False

    for line in md_content.split('\n'):
        stripped = line.strip()

        # Empty lines
        if not stripped:
            if in_bibliography:
                continue
            renderer.add_empty_line()
            continue

        # Separator → bibliografia title
        if stripped == '---' and not in_bibliography:
            in_bibliography = True
            renderer.add_empty_line()
            renderer.add_bibliography_title()
            renderer.add_empty_line()
            continue

        # Headings
        if stripped.startswith('# '):
            text = stripped[2:].strip()
            renderer.add_chapter_title(text)
        elif stripped.startswith('## '):
            text = stripped[3:].strip()
            m = re.match(r'^(\d+)\.\s+(.+)$', text)
            if m:
                renderer.add_section(m.group(1), m.group(2))
            else:
                renderer.add_section_heading(text)
        elif stripped.startswith('### '):
            renderer.add_author(stripped[4:].strip())
        elif stripped.startswith('#### '):
            text = stripped[5:].strip()
            # Detect section numbering and split number from title
            m = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', text)
            if m:
                num, title = m.group(1), m.group(2)
                if re.match(r'^\d+\.\d+\.\d+$', num):
                    renderer.add_subsubsection(num, title)
                elif re.match(r'^\d+\.\d+$', num):
                    renderer.add_subsection(num, title)
                else:
                    renderer.add_section(num, title)
            else:
                renderer.add_raw_paragraph(text, italic=True, alignment=ALIGN_JUSTIFY)

        # Bibliography entries
        elif in_bibliography and stripped.startswith('- '):
            text = stripped[2:]
            # Detect [titolo in corsivo] pattern
            match = re.search(r'\[([^\]]+)\]', text)
            if match:
                before = text[:match.start()]
                title = match.group(1)
                after = text[match.end():]
                renderer.add_bibliography_entry_parts(
                    (before, False),
                    (title, True),
                    (after, False),
                )
            else:
                renderer.add_bibliography_entry(text)

        # Bold = figure caption
        elif stripped.startswith('**') and stripped.endswith('**'):
            renderer.add_figure_caption(stripped[2:-2])

        # Regular paragraph
        else:
            renderer.add_paragraph(stripped)

    return renderer


def cli():
    parser = argparse.ArgumentParser(
        description="Converti Markdown in DOCX formattato Roma TrE-PRESS"
    )
    parser.add_argument("input", help="File Markdown di input")
    parser.add_argument("output", help="File DOCX di output")
    parser.add_argument(
        "--template",
        default="raw/articles/template_Psicology & Education.docx",
        help="Percorso del template DOCX (default: raw/articles/template_Psicology & Education.docx)",
    )
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        md_content = f.read()

    renderer = markdown_to_docx(md_content, args.template)
    renderer.save(args.output)
    print(f"✅ Documento salvato: {args.output}")


if __name__ == "__main__":
    cli()
